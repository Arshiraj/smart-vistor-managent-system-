import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_text_formatting(cell, font_size=9.5, bold=False, rgb_color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if rgb_color:
                run.font.color.rgb = rgb_color

# Try to load existing word file
input_docx = "SkillMapper AI_IDF_New Submission_25-Feb-2026.docx"
output_docx = "Project_Jyoti_Final_With_TestCases.docx"

try:
    doc = Document(input_docx)
    doc.add_page_break()
    doc.add_heading('Appendix: Jyoti Core Test Cases', level=1)
except Exception as e:
    # If the file fails to open, just create a new one
    print("Could not open existing document, creating a standalone one: ", e)
    doc = Document()
    doc.add_heading('Appendix: Jyoti Core Test Cases', level=1)

headers = ['TC #', 'Module', 'Test Description', 'Input', 'Expected', 'Actual', 'Status']
data = [
    ['TC01', 'Biometrics', 'Successful Biometric Scan & Tokenization', 'Registered face scan', 'Token generated; ACCESS GRANTED', 'Token generated', '✓ Pass'],
    ['TC02', 'Biometrics', 'Liveness Detection Rejection', 'Static photo scan', 'Access denied; Spoof Attempt alert', 'Access denied', '✓ Pass'],
    ['TC03', 'Zone Control', 'Authorized Zone Transition', 'Move Reception -> Zone-A', 'ZONE_CHANGE event; No alerts', 'Heatmap updated', '✓ Pass'],
    ['TC04', 'Zone Control', 'Wandering Visitor Detection', 'Enter unauthorized Server Room', 'CRITICAL alert; UI flashes red', 'Alert triggered', '✓ Pass'],
    ['TC05', 'Zone Control', 'Passback Violation', '2 rapid distant token scans', 'Passback Anomaly; Suspended', 'Token suspended', '✓ Pass'],
    ['TC06', 'Threat Detect', 'Normal Sentiment Processing', 'thermal: 36.6, sentiment: Calm', 'Continues movement loop', 'Loop continued', '✓ Pass'],
    ['TC07', 'Threat Detect', 'High-Agitation Behavioral Focus', 'Sentiment Agitated for 10s', 'Marked flagged: true, WARNING', 'Indicator red', '✓ Pass'],
    ['TC08', 'Threat Detect', 'Thermal Escalation Trigger', 'Core temp reading 38.5', 'CRITICAL threat alert', 'Action required', '✓ Pass'],
    ['TC09', 'System Sync', 'Live Settings Deployment', 'Adjusted sensitivity slider', 'Configs apply to securityEngine', 'Sens updated', '✓ Pass'],
    ['TC10', 'System Sync', 'Immediate Incident Broadcast', 'Guard broadcasts Evacuate', 'Global DANGER banner', 'Banner displayed', '✓ Pass'],
    ['TC11', 'Escalation', 'Incident Resolution Cycle', 'Acknowledge -> Generate PDF', 'ACKNOWLEDGED state -> report', 'Report generated', '✓ Pass'],
    ['TC12', 'Audit Logs', 'Audit Log Integration', 'Diverse safe & flagged events', 'Data mapped securely', 'Data filtered', '✓ Pass']
]

table = doc.add_table(rows=1, cols=7)
table.style = 'Normal Table'

hdr_cells = table.rows[0].cells
header_hex = '1D3C5A'
header_rgb = RGBColor(255, 255, 255)

for i, title in enumerate(headers):
    hdr_cells[i].text = title
    set_cell_background(hdr_cells[i], header_hex)
    set_text_formatting(hdr_cells[i], font_size=10, bold=True, rgb_color=header_rgb)

dark_blue_text = RGBColor(29, 60, 90)
cream_bg = 'F5EFE6'
white_bg = 'FFFFFF'
black_text = RGBColor(0, 0, 0)

for index, row_data in enumerate(data):
    row_cells = table.add_row().cells
    bg_color = cream_bg if index % 2 == 0 else white_bg
    for i, text in enumerate(row_data):
        row_cells[i].text = text
        set_cell_background(row_cells[i], bg_color)
        
        if i == 0:
            set_text_formatting(row_cells[i], font_size=9.5, bold=True, rgb_color=dark_blue_text)
        else:
            set_text_formatting(row_cells[i], font_size=9.5, bold=False, rgb_color=black_text)

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

doc.save(output_docx)
print(f"Successfully inserted test cases and saved to {output_docx}")
